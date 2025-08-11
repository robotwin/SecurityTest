import os
from docx import Document
from openpyxl import Workbook
#Copyright (C) 2025 Carizon12345
#获取待处理的文件的路径
#hello88
path='word文件'  #文件所在文件夹
files = [path+"\\"+i for i in os.listdir(path)] #获取文件夹下的文件名,并拼接完整路径

for file in files:
    doc = Document(file)

    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中
            
    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")
    wb = Workbook()
    wb.remove(wb.worksheets[0])  # 删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1):  # 从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")  # 创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)):  # 遍历word中表格的所有行
            row_data = []  # 储存表格中每行的数据
            for j in range(len(table.columns)):  # 遍历word中表格的所有列
                row_data.append(table.cell(i, j).text)
            ws.append(row_data)  # 每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0]))  # 保存excel文件
    # commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")
    wb = Workbook()
    wb.remove(wb.worksheets[0])  # 删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1):  # 从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")  # 创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)):  # 遍历word中表格的所有行
            row_data = []  # 储存表格中每行的数据
            for j in range(len(table.columns)):  # 遍历word中表格的所有列
                row_data.append(table.cell(i, j).text)
            ws.append(row_data)  # 每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0]))  # 保存excel文件
    # commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")    wb = Workbook()
    wb.remove(wb.worksheets[0])#删除工作簿自带的工作表
    for index, table in enumerate(doc.tables, start=1): #从1开始给表格编号
        ws = wb.create_sheet(f"Sheet{index}")#创建新工作表，以"Sheet" + word中表格的编号命名
        for i in range(len(table.rows)): #遍历word中表格的所有行
            row_data = [] #储存表格中每行的数据
            for j in range(len(table.columns)): #遍历word中表格的所有列
                row_data.append(table.cell(i,j).text)
            ws.append(row_data) #每取一行就写入数据到Excel表的行中

    wb.save("excel文件\\{}.xlsx".format(file.split("\\")[1].split(".")[0])) #保存excel文件
    #commit the change
    print("Hello this code is a pluuidwhfoiufhalksdjfhalskdjhfs")













































































